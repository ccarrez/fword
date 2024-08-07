import os
from docx import Document
from odf import text, teletype
from odf.opendocument import load

oldText = input("Entrer le texte à chercher : ")
newText = input("Entrer le texte à remplacer : ")

# search and replace a text in a docx document
files_src = [name for name in os.listdir('./todo') if name.endswith(".docx")]

for file_src in files_src:
    
    # log
    print("Processing " + file_src)

    # load input document
    document = Document('./todo/'+file_src)

    # search and replace in paragraphs
    for paragraph in document.paragraphs:
        if oldText in paragraph.text:
            if paragraph.text == '': continue
            paragraph.text = paragraph.text.replace(oldText, newText)

    # search and replace in tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text == '': continue
                    paragraph.text = paragraph.text.replace(oldText, newText)

    # search and replace in header
    section = document.sections[0]
    header = section.header
    for paragraph in header.paragraphs:
        if oldText in paragraph.text:
            if paragraph.text == '': continue
            paragraph.text = paragraph.text.replace(oldText, newText)
    
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text == '': continue
                    paragraph.text = paragraph.text.replace(oldText, newText)

    # search and replace in footer
    section = document.sections[len(document.sections)-1]
    footer = section.footer
    for paragraph in footer.paragraphs:
        if oldText in paragraph.text:
            if paragraph.text == '': continue
            paragraph.text = paragraph.text.replace(oldText, newText)
    
    for table in footer.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text == '': continue
                    paragraph.text = paragraph.text.replace(oldText, newText)

    # save modified document in another folder
    document.save('./done/'+file_src)  

# search and replace a text in a odt document
files_src = [name for name in os.listdir('./todo') if name.endswith(".odt")]

for file_src in files_src:
    # log
    print("Processing " + file_src)

    # load input document
    document = load('./todo/'+file_src)

    # search and replace in paragraphs
    paragraphs = document.getElementsByType(text.P)
    for paragraph in paragraphs:        
        old_text = teletype.extractText(paragraph)
        if old_text == '': continue
        new_text = old_text.replace(oldText, newText)
        new_S = text.P()
        new_S.setAttribute("stylename", paragraph.getAttribute("stylename"))
        new_S.addText(new_text)
        paragraph.parentNode.insertBefore(new_S, paragraph)
        paragraph.parentNode.removeChild(paragraph)

    document.save('./done/'+file_src)